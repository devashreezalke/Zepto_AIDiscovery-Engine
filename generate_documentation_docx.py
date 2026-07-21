import sys
import subprocess
import os

# Auto-install python-docx if not present
try:
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import nsdecls, qn
except ImportError:
    print("python-docx not found. Installing now...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    """Sets background shading of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def add_heading_with_spacing(doc, text, level, space_before=12, space_after=6, color_rgb=RGBColor(62, 0, 117)):
    heading = doc.add_heading(text, level=level)
    heading.paragraph_format.space_before = Pt(space_before)
    heading.paragraph_format.space_after = Pt(space_after)
    for run in heading.runs:
        run.font.name = 'Arial'
        run.font.color.rgb = color_rgb
    return heading

def add_paragraph_with_spacing(doc, text, space_before=0, space_after=6, italic=False, bold=False, color_rgb=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.font.italic = italic
    run.font.bold = bold
    if color_rgb:
        run.font.color.rgb = color_rgb
    return p

def create_system_documentation():
    doc = docx.Document()
    
    # Page setup
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    
    # Title
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(2)
    run_title = p_title.add_run("Zepto AI Discovery Engine: System Documentation")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(62, 0, 117)
    
    # Subtitle
    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(24)
    run_sub = p_sub.add_run("Technical overview of data gathering, NLP enrichment, DBSCAN clustering, LLM synthesis prompts, and validation gate mathematics.")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(12)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(107, 114, 128)

    # 1. Section 1: Ingestion
    add_heading_with_spacing(doc, "1. How the Workflow Gathers and Analyzes Data", level=1, space_before=18)
    
    add_paragraph_with_spacing(doc, "Data Gathering & Ingestion Layer", bold=True, color_rgb=RGBColor(62, 0, 117))
    add_paragraph_with_spacing(doc, "The Zepto AI Discovery Engine operates a multi-platform ingestion workflow that automatically consolidates raw user feedback from 4 sources:")
    
    bullets_sources = [
        "Play Store Connector: Extracts reviews for 'com.zeptoconsumerapp' and stores ratings, date, and version details.",
        "Reddit API Crawler: Queries subreddits like r/zepto, r/bangalore, and r/india for user discussions regarding logistics and customer experiences.",
        "Twitter Snscrape: Queries tweets matching quick-commerce discovery friction terms.",
        "Forums Scraper: Parses consumer complaint forums and product blogs using BeautifulSoup to extract comparative customer feedback."
    ]
    for bullet in bullets_sources:
        p_b = doc.add_paragraph(style='List Bullet')
        p_b.paragraph_format.space_after = Pt(4)
        run_b = p_b.add_run(bullet)
        run_b.font.name = 'Arial'
        run_b.font.size = Pt(10.5)

    add_paragraph_with_spacing(doc, "PII Cleansing and Exact Deduplication", bold=True, space_before=8)
    add_paragraph_with_spacing(doc, "To maintain data compliance and security, all raw text records are processed through regex filters to redact phone numbers, email addresses, and names. Additionally, to preserve LLM API limits, the ingestion layer compares and flags exact duplicates (setting is_duplicate = 1 in SQLite) to bypass duplicate LLM calls while preserving total occurrence counts for statistics.")

    add_paragraph_with_spacing(doc, "NLP & LLM Enrichment Analysis", bold=True, space_before=8)
    add_paragraph_with_spacing(doc, "The analyzer engine reads unique documents and queries the Groq API (using llama-3.1-8b-instant for high-throughput). The prompt forces the LLM to output a clean JSON schema including:")
    
    bullets_nlp = [
        "Sentiment Score: A bounded decimal between -1.0 (highly negative friction) and +1.0 (positive trial experience).",
        "Target Blocked Categories: Identifies which specific shopping folder was blocked (e.g. Groceries, Baby Care, Electronics).",
        "Primary Customer Drivers: Mapped to categories like trust deficit, price comparison, muscle-memory habit, or logistics delays.",
        "Research Question Mapping: Tagged against Q1-Q8 codes based on matching content markers."
    ]
    for bullet in bullets_nlp:
        p_b = doc.add_paragraph(style='List Bullet')
        p_b.paragraph_format.space_after = Pt(4)
        run_b = p_b.add_run(bullet)
        run_b.font.name = 'Arial'
        run_b.font.size = Pt(10.5)

    # 2. Section 2: Themes
    add_heading_with_spacing(doc, "2. How Themes Are Identified", level=1, space_before=18)
    add_paragraph_with_spacing(doc, "The engine uses a density-based spatial clustering approach to group semantically similar customer complaints without manual categorizing:")
    
    add_paragraph_with_spacing(doc, "Vector Embeddings Generation", bold=True, color_rgb=RGBColor(62, 0, 117))
    add_paragraph_with_spacing(doc, "Text entries are converted into 384-dimensional dense vector embeddings using the sentence-transformers model. These embeddings map semantic similarities in a coordinate space where physically close vectors represent complaints discussing similar issues (even when using different vocabulary). To handle memory limits on standard machines, the engine implements a category-driver coordinate system that dynamically maps text categories into orthogonal vector space.")
    
    add_paragraph_with_spacing(doc, "DBSCAN Density Clustering", bold=True, space_before=8)
    add_paragraph_with_spacing(doc, "The vector space is clustered using DBSCAN (Density-Based Spatial Clustering of Applications with Noise). Epsilon and min_samples parameters are set to group dense centers while filtering isolated feedback as background noise:")
    
    p_eq1 = doc.add_paragraph()
    p_eq1.paragraph_format.left_indent = Inches(0.5)
    run_eq1 = p_eq1.add_run("Group criteria: dist(x_i, x_j) < Epsilon  |  Cluster size >= MinSamples")
    run_eq1.font.name = 'Courier New'
    run_eq1.font.size = Pt(10.5)
    run_eq1.font.bold = True
    run_eq1.font.color.rgb = RGBColor(226, 26, 132)

    add_paragraph_with_spacing(doc, "Dimensionality Reduction & PCA Plotting", bold=True, space_before=8)
    add_paragraph_with_spacing(doc, "For visualization, the high-dimensional coordinates are projected onto a 2D plane using Principal Component Analysis (PCA). The resulting scatter plot is saved in the data folder, illustrating cluster separation. Centroid reviews are evaluated, and the cluster is labeled with a descriptive macro-theme (e.g. 'Quality and Trust barriers for Groceries').")

    # 3. Section 3: Insights
    add_heading_with_spacing(doc, "3. How Insights Are Generated", level=1, space_before=18)
    add_paragraph_with_spacing(doc, "Once theme clusters are established, the engine synthesizes them into actionable growth product initiatives using a structured PM framework:")
    
    add_paragraph_with_spacing(doc, "Growth PM Synthesis Prompt", bold=True, color_rgb=RGBColor(62, 0, 117))
    add_paragraph_with_spacing(doc, "For each theme with significant volume, the system gathers its centroid summary, document count, average sentiment score, and representative customer quotes. It formats these metrics into a synthesis prompt sent to the LLM (llama-3.3-70b-versatile).")
    
    add_paragraph_with_spacing(doc, "If-Then-Leading-To Hypothesis Structure", bold=True, space_before=8)
    add_paragraph_with_spacing(doc, "The model is constrained to formulate growth experiments strictly using the following PM-standard format:")
    
    p_eq2 = doc.add_paragraph()
    p_eq2.paragraph_format.left_indent = Inches(0.5)
    p_eq2.paragraph_format.space_before = Pt(4)
    p_eq2.paragraph_format.space_after = Pt(4)
    run_eq2 = p_eq2.add_run("If we [UX / Operational Intervention], \nThen [User Reaction / Action], \nLeading to [Target Metric Lift].")
    run_eq2.font.name = 'Arial'
    run_eq2.font.size = Pt(11)
    run_eq2.font.italic = True
    run_eq2.font.bold = True
    run_eq2.font.color.rgb = RGBColor(62, 0, 117)

    add_paragraph_with_spacing(doc, "For each hypothesis, the model must also specify the experiment type (e.g. A/B test), required engineering effort (Low/Medium/High), and expected business impact.")

    # 4. Section 4: Validation
    add_heading_with_spacing(doc, "4. How We Validated the Quality of the Insights", level=1, space_before=18)
    add_paragraph_with_spacing(doc, "To prevent product teams from pursuing false trends, the engine runs every insight through a 5-Gate Validation Framework that computes a weighted composite confidence score:")
    
    add_paragraph_with_spacing(doc, "The 5 Quality Gates", bold=True, color_rgb=RGBColor(62, 0, 117))
    
    gates = [
        "Gate 1 (Volume): Ensures the insight is supported by a large cluster. (Weight: 0.25)",
        "Gate 2 (Diversity): Verifies the issue is mentioned across multiple data sources (e.g. App Store + Reddit). (Weight: 0.20)",
        "Gate 3 (Consistency): Analyzes the standard deviation of sentiment scores within the cluster to ensure consistent customer feedback. (Weight: 0.20)",
        "Gate 4 (Temporal Span): Evaluates the time duration of complaints to verify it's a persistent issue rather than a temporary server glitch. (Weight: 0.15)",
        "Gate 5 (LLM Self-Consistency): Checks model agreement checks. (Weight: 0.20)"
    ]
    for gate in gates:
        p_g = doc.add_paragraph(style='List Bullet')
        p_g.paragraph_format.space_after = Pt(4)
        run_g = p_g.add_run(gate)
        run_g.font.name = 'Arial'
        run_g.font.size = Pt(10.5)

    add_paragraph_with_spacing(doc, "Mathematical Confidence Score Formula", bold=True, space_before=8)
    add_paragraph_with_spacing(doc, "The composite score is calculated using the following algebraic formula:")
    
    p_eq3 = doc.add_paragraph()
    p_eq3.paragraph_format.left_indent = Inches(0.5)
    run_eq3 = p_eq3.add_run("Confidence = (Vol * 0.25) + (Div * 0.20) + (Cons * 0.20) + (Temp * 0.15) + (LLM * 0.20)")
    run_eq3.font.name = 'Courier New'
    run_eq3.font.size = Pt(11)
    run_eq3.font.bold = True
    run_eq3.font.color.rgb = RGBColor(226, 26, 132)

    add_paragraph_with_spacing(doc, "Validation Thresholds", bold=True, space_before=8)
    add_paragraph_with_spacing(doc, "Based on the resulting composite score, the validation engine classifies the insight's confidence level:")
    
    p_t1 = doc.add_paragraph()
    p_t1.paragraph_format.left_indent = Inches(0.4)
    p_t1.paragraph_format.space_after = Pt(2)
    p_t1.add_run("• Score >= 0.80: ").font.bold = True
    p_t1.add_run("HIGH Confidence (Green Light - ready to design sprint A/B test)")
    p_t1.runs[0].font.name = 'Arial'
    p_t1.runs[1].font.name = 'Arial'
    
    p_t2 = doc.add_paragraph()
    p_t2.paragraph_format.left_indent = Inches(0.4)
    p_t2.paragraph_format.space_after = Pt(2)
    p_t2.add_run("• 0.60 <= Score < 0.80: ").font.bold = True
    p_t2.add_run("MEDIUM Confidence (Orange Light - request qualitative follow-ups)")
    p_t2.runs[0].font.name = 'Arial'
    p_t2.runs[1].font.name = 'Arial'
    
    p_t3 = doc.add_paragraph()
    p_t3.paragraph_format.left_indent = Inches(0.4)
    p_t3.paragraph_format.space_after = Pt(12)
    p_t3.add_run("• Score < 0.60: ").font.bold = True
    p_t3.add_run("LOW Confidence (Red Light - discard or monitor for future feedback)")
    p_t3.runs[0].font.name = 'Arial'
    p_t3.runs[1].font.name = 'Arial'
    
    # Save document
    filename = "Zepto_AI_Discovery_Engine_System_Documentation.docx"
    doc.save(filename)
    print(f"System documentation saved successfully as '{os.path.abspath(filename)}'")

if __name__ == '__main__':
    create_system_documentation()
