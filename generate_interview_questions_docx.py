import sys
import subprocess
import os

# Auto-install python-docx if not present
try:
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import nsdecls, qn
except ImportError:
    print("python-docx not found. Installing now...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    """Sets background shading of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def add_heading_with_spacing(doc, text, level, space_before=12, space_after=6, color_rgb=RGBColor(62, 0, 117)):
    heading = doc.add_heading(text, level=level)
    heading.paragraph_format.space_before = Pt(space_before)
    heading.paragraph_format.space_after = Pt(space_after)
    for run in heading.runs:
        run.font.name = 'Arial'
        run.font.color.rgb = color_rgb
    return heading

def add_paragraph_with_spacing(doc, text, space_before=0, space_after=6, italic=False, bold=False, color_rgb=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.font.italic = italic
    run.font.bold = bold
    if color_rgb:
        run.font.color.rgb = color_rgb
    return p

def create_interview_guide():
    doc = docx.Document()
    
    # Page setup
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    
    # Title
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(2)
    run_title = p_title.add_run("User Interview Guide: High-Value Explorers")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(20)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(62, 0, 117)
    
    # Subtitle
    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(24)
    run_sub = p_sub.add_run("Qualitative interview script to uncover trust barriers, purchase partition habits, and validation needs for non-grocery shopping on Zepto.")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(11)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(107, 114, 128)

    # Intro Guidelines
    add_heading_with_spacing(doc, "Interview Guidelines & Best Practices", level=1, space_before=14)
    
    p_intro = doc.add_paragraph()
    p_intro.paragraph_format.space_after = Pt(4)
    run_i1 = p_intro.add_run("• Who to Interview: ")
    run_i1.font.bold = True
    run_i2 = p_intro.add_run("Dual Income No Kids (DINKs) couples, urban millennial parents, and pet owners (age 22–40) in Tier-1 cities who order groceries on Zepto but buy beauty, electronics, baby, or pet care elsewhere.")
    run_i1.font.name = 'Arial'
    run_i2.font.name = 'Arial'
    
    p_format = doc.add_paragraph()
    p_format.paragraph_format.space_after = Pt(4)
    run_f1 = p_format.add_run("• Interview Format: ")
    run_f1.font.bold = True
    run_f2 = p_format.add_run("30-minute semi-structured video calls with screen-sharing active. Ask users to show their recent order history on Zepto and competitor apps.")
    run_f1.font.name = 'Arial'
    run_f2.font.name = 'Arial'
    
    p_tone = doc.add_paragraph()
    p_tone.paragraph_format.space_after = Pt(12)
    run_t1 = p_tone.add_run("• Interviewer Tone: ")
    run_t1.font.bold = True
    run_t2 = p_tone.add_run("Non-judgmental, curious. Focus on concrete past behaviors and stories rather than hypothetical user preferences.")
    run_t1.font.name = 'Arial'
    run_t2.font.name = 'Arial'

    # Section 1
    add_heading_with_spacing(doc, "Section 1: App Usage Context & Warm-Up", level=2, space_before=14, color_rgb=RGBColor(226, 26, 132))
    questions_s1 = [
        "Can you walk me through the last time you opened Zepto? What did you search for and what did you buy?",
        "When you think of the name 'Zepto', what is the first category of products that comes to your mind?",
        "Besides groceries and fresh produce, what other folders or items have you noticed on the app's homepage?"
    ]
    for q in questions_s1:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(q)
        run.font.name = 'Arial'
        run.font.size = Pt(10.5)

    # Section 2
    add_heading_with_spacing(doc, "Section 2: Grocery Muscle Memory & Partitioning", level=2, space_before=14)
    questions_s2 = [
        "How do you decide which shopping app to open? For instance, what triggers you to open Zepto vs. Amazon vs. Blinkit vs. Nykaa?",
        "Have you ever bought something like a charging cable, a toy for your pet, or a baby soap on Zepto? If yes, what was that experience like? If no, what made you go elsewhere?",
        "When buying non-grocery items (e.g. skin creams, baby products), how do you typically buy them? Describe your last purchase."
    ]
    for q in questions_s2:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(q)
        run.font.name = 'Arial'
        run.font.size = Pt(10.5)

    # Section 3
    add_heading_with_spacing(doc, "Section 3: Catalog Browsing & Cart Drop-off", level=2, space_before=14, color_rgb=RGBColor(226, 26, 132))
    questions_s3 = [
        "Have you ever browsed through folders like 'Beauty & Cosmetics' or 'Electronics' on Zepto out of curiosity, added an item to your cart, but decided not to checkout? Can you recall what stopped you?",
        "When looking at an unfamiliar brand or item on Zepto, what details do you look for on the product page before clicking 'Add to Cart'?",
        "How do you feel about the quality of items that Zepto delivery boys select for you? Does that affect your decision to buy baby care or premium cosmetics?"
    ]
    for q in questions_s3:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(q)
        run.font.name = 'Arial'
        run.font.size = Pt(10.5)

    # Section 4
    add_heading_with_spacing(doc, "Section 4: Trust, Specifications & Returns Anxiety", level=2, space_before=14)
    questions_s4 = [
        "If you receive a defective electronic item (e.g. a keyboard) or a broken cosmetics bottle, what is your expectation of the return process? Have you ever had to deal with Zepto customer support for a refund?",
        "When shopping for baby diapers or pet collars, what specific sizing or fitting information do you need? Where do you usually find this information today?",
        "How important are user reviews and star ratings to you when trying a new snack or personal care item? How do you feel about Zepto not showing customer ratings on product listings?"
    ]
    for q in questions_s4:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(q)
        run.font.name = 'Arial'
        run.font.size = Pt(10.5)

    # Section 5
    add_heading_with_spacing(doc, "Section 5: Reactivity to Proposed Solutions", level=2, space_before=14, color_rgb=RGBColor(226, 26, 132))
    questions_s5 = [
        "If Zepto added star ratings and customer reviews (similar to Amazon/Blinkit) directly on all cosmetics and snacks PDPs, how would that change your confidence in buying them?",
        "If you saw a prominent 'No-Questions-Asked Instant Refund' badge on premium electronics, how would that impact your hesitation to order a device under 10 minutes?",
        "If diaper and baby care listings featured clear weight and sizing tables directly on the page, would you start ordering these items on Zepto instead of dedicated parenting apps?"
    ]
    for q in questions_s5:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(q)
        run.font.name = 'Arial'
        run.font.size = Pt(10.5)

    # Save document
    filename = "Zepto_High_Value_Explorers_Interview_Guide.docx"
    doc.save(filename)
    print(f"Interview guide saved successfully as '{os.path.abspath(filename)}'")

if __name__ == '__main__':
    create_interview_guide()
