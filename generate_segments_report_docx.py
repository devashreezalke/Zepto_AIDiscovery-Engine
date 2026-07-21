import sys
import subprocess
import os

# Auto-install python-docx if not present
try:
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import nsdecls, qn
except ImportError:
    print("python-docx not found. Installing now...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    """Sets background shading of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def add_heading_with_spacing(doc, text, level, space_before=12, space_after=6, color_rgb=RGBColor(62, 0, 117)):
    heading = doc.add_heading(text, level=level)
    heading.paragraph_format.space_before = Pt(space_before)
    heading.paragraph_format.space_after = Pt(space_after)
    for run in heading.runs:
        run.font.name = 'Arial'
        run.font.color.rgb = color_rgb
    return heading

def add_paragraph_with_spacing(doc, text, space_before=0, space_after=6, italic=False, bold=False, color_rgb=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.font.italic = italic
    run.font.bold = bold
    if color_rgb:
        run.font.color.rgb = color_rgb
    return p

def create_combined_segments_report():
    doc = docx.Document()
    
    # Page setup
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    
    # Title
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(2)
    run_title = p_title.add_run("Zepto Target Customer Strategy Report")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(62, 0, 117)
    
    # Subtitle
    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(24)
    run_sub = p_sub.add_run("Focusing on the combined High-Value Explorers segment to drive quick-commerce margin expansion, basket growth, and active cross-category conversion.")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(11)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(107, 114, 128)

    # 1. Executive Segmentation Matrix
    add_heading_with_spacing(doc, "1. Executive Segmentation Matrix", level=1, space_before=18)
    
    headers = ["Segment Name", "Size (%)", "Target Priority", "Demographic Focus", "Primary Growth Action"]
    table_data = [
        ["High-Value Explorers (Combined)", "40%", "NORTH STAR FOCUS", "Millennials, DINKs, young parents, pet owners", "PDP star reviews & trust badges"],
        ["Habitual Minimalists", "30%", "Secondary Focus", "Single professionals, 20-30", "Post-cart cross-sells"],
        ["Value Seekers", "20%", "Low Priority", "Middle-class families, 28-45", "Bulk packs & fee waivers"],
        ["Platform Loyalists", "10%", "Medium Priority", "High-income Zepto Pass users", "Early access category launches"]
    ]
    
    table = doc.add_table(rows=1, cols=5)
    table.autofit = True
    
    hdr_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        hdr_cells[i].text = header_text
        set_cell_background(hdr_cells[i], "3E0075")
        p = hdr_cells[i].paragraphs[0]
        p.runs[0].font.bold = True
        p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
        p.runs[0].font.name = 'Arial'
        p.runs[0].font.size = Pt(9.5)
        
    for row_idx, data in enumerate(table_data):
        row_cells = table.add_row().cells
        for col_idx, text in enumerate(data):
            row_cells[col_idx].text = text
            p = row_cells[col_idx].paragraphs[0]
            p.runs[0].font.name = 'Arial'
            p.runs[0].font.size = Pt(9)
            if row_idx % 2 == 1:
                set_cell_background(row_cells[col_idx], "F9FAFB")
            else:
                set_cell_background(row_cells[col_idx], "FFFFFF")

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # 2. Detailed Profiles
    add_heading_with_spacing(doc, "2. Deep-Dive User Segment Profiles", level=1, space_before=18)
    
    segments_detailed = [
        {
            "name": "High-Value Explorers (40% of MACs) — COMBINED NORTH STAR FOCUS",
            "color": RGBColor(226, 26, 132), # Magenta
            "demographics": "• Age Cohort: 22–40 years old.\n• Household Profile: Dual Income No Kids (DINKs) couples, young parents with infants/toddlers (0-5 years), and active pet owners. Living in high-density premium apartments.\n• Income Bracket: High disposable income. Willing to spend more on organic wellness, premium cosmetics, baby health, or quality electronics.\n• Location: Prime residential pockets and IT corridors of Tier-1 metros (Bangalore, Mumbai, Delhi NCR, Hyderabad).",
            "behavioral": "• Session Duration: Long to Medium (averaging 4–10 minutes per session).\n• Navigation Pattern: Deep scrolling through folder hierarchies, views adjacent categories (Beauty, Organic, Baby care), actively queries specific brands, and reads PDP ingredient/specification details.\n• Purchase Frequency: High (2–3 times per week).\n• Basket Characteristics: Premium Average Order Value (AOV ~600+ INR) with heavy indexing on diapers, premium skin care, pet treats, and organic vegetables.",
            "barriers": "• Quality & Authenticity Trust Deficit: Skepticism regarding picker boys selecting bruised produce, and fear of unauthentic premium cosmetic/pet brands.\n• Information Gaps: Lack of star ratings or customer feedback reviews on PDPs to guide purchases, and missing sizing/weight grids for infant diapers.\n• Return/Refund Opaque Rules: Fear of unhelpful scripted support chatbots rejecting refund requests for missing/defective goods.",
            "intervention": "Integrate star ratings & review snippets on all non-grocery PDPs, add diaper sizing weight grids, and introduce 'No-Questions-Asked Instant Refund' and '100% Brand Certified' trust badges on category listings."
        },
        {
            "name": "Habitual Minimalists (30% of MACs)",
            "color": RGBColor(62, 0, 117), # Purple
            "demographics": "• Age Cohort: 20–30 years old.\n• Household Profile: Single working professionals, students, or shared-apartment dwellers.\n• Income Bracket: Lower-Middle to Middle.\n• Location: Co-living complexes in major cities.",
            "behavioral": "• Session Duration: Ultra-short (averaging < 120 seconds).\n• Navigation Pattern: Bypasses the folder catalog entirely. Relies strictly on search bar autocomplete, order history, or direct cart checkouts.\n• Purchase Frequency: High (3–4 times per week, emergency-driven).\n• Basket Characteristics: Small AOV (~200 INR), consisting of 2-3 immediate staples (milk, bread, onions).",
            "barriers": "• High Cognitive Load: Explores no banners because checkout is a single-focus, transactional task.\n• Pricing Sensitivity: Highly sensitive to delivery, handling, and surge convenience fee markups.",
            "intervention": "Inject contextual cross-sells at cart checkout for staples (e.g. recommend fresh coriander/chilies when buying tomatoes and onions) to capture adjacent cross-buying without adding browsing friction."
        },
        {
            "name": "Value Seekers (20% of MACs)",
            "color": RGBColor(107, 114, 128),
            "demographics": "• Age Cohort: 28–45 years old.\n• Household Profile: Budget-conscious family households.\n• Income Bracket: Middle.\n• Location: Tier-1 and Tier-2 urban residential blocks.",
            "behavioral": "• Session Duration: Medium to Long (4–8 minutes).\n• Navigation Pattern: Immediately visits the 'Offers' or 'Deals' tab. Searches for discount bundles.\n• Purchase Frequency: Moderate (2-3 times per month, bulk buys).\n• Basket Characteristics: Large items-per-basket count, highly sensitive to handling, packing, and surge fee fluctuations.",
            "barriers": "• Premium Price Perception: Compares price-per-kg of vegetables and groceries with local bulk supermarkets (e.g. D-Mart) and assumes quick-commerce is strictly a premium convenience option.",
            "intervention": "Promote bulk-saver deals, display price-per-kg comparison tags on produce, and set threshold indicators to waive delivery/surge charges."
        },
        {
            "name": "Platform Loyalists (10% of MACs)",
            "color": RGBColor(107, 114, 128),
            "demographics": "• Age Cohort: 25–50 years old.\n• Household Profile: Tech-savvy professionals and premium quick-commerce adapters.\n• Income Bracket: High.\n• Location: High-value residential neighborhoods.",
            "behavioral": "• Session Duration: Medium (3–4 minutes).\n• Navigation Pattern: Navigates across multiple folders. Uses Zepto Pass premium subscription benefits.\n• Purchase Frequency: Very High (> 10 orders per month).\n• Basket Characteristics: High AOV (>600 INR), already purchases from 3+ distinct categories monthly.",
            "barriers": "• Out-of-Stock Friction: Customer satisfaction drops immediately if peak-time slots are delayed or favorite staple items are missing.",
            "intervention": "Provide VIP early access to new category folder launches and priority customer care channels to maintain high brand loyalty."
        }
    ]

    for seg in segments_detailed:
        add_heading_with_spacing(doc, seg["name"], level=2, space_before=14, space_after=4, color_rgb=seg["color"])
        
        # Demographics
        add_paragraph_with_spacing(doc, "Demographic Profile:", bold=True, space_after=1, color_rgb=seg["color"])
        add_paragraph_with_spacing(doc, seg["demographics"], space_after=6)
        
        # Behavioral
        add_paragraph_with_spacing(doc, "Behavioral Characteristics:", bold=True, space_after=1)
        add_paragraph_with_spacing(doc, seg["behavioral"], space_after=6)
        
        # Friction/Barriers
        add_paragraph_with_spacing(doc, "Friction & Exploration Barriers:", bold=True, space_after=1)
        add_paragraph_with_spacing(doc, seg["barriers"], space_after=6)
        
        # Custom trigger
        add_paragraph_with_spacing(doc, "Recommended Product Intervention:", bold=True, space_after=1, color_rgb=RGBColor(62, 0, 117))
        add_paragraph_with_spacing(doc, seg["intervention"], space_after=14, italic=True)

    # 3. Focus Rationale
    add_heading_with_spacing(doc, "3. Rationale: Why Focus on High-Value Explorers?", level=1, space_before=18)
    
    rationales = [
        ("1. Maximum Contribution to Gross Margin", 
         "Staples like milk, bread, and onions yield tiny margins (<5%). High-Value Explorers buy from folders that yield premium quick-commerce gross margins of 18–25% (organic oil, baby toiletries, cosmetics, pet toys, charging cables). Activating this segment shifts catalog mix directly to higher profitability."),
        ("2. Active Exploration Intent is Already Unlocked", 
         "This segment spends 4–10 minutes browsing, exploring PDPs, and searching for premium brands. They have intent; they fail to checkout strictly due to trust, specification, and risk barriers. Resolving these barriers converts existing traffic without new acquisition costs."),
        ("3. Massive Target Scale (40% of MACs)", 
         "Combining Curious Browsers (25%) and Life-Stage Movers (15%) creates a dominant segment representing 40% of Zepto's monthly active traffic. Converting this single segment delivers massive scale for catalog penetration and AOV expansion.")
    ]
    for title, desc in rationales:
        add_paragraph_with_spacing(doc, title, bold=True, space_after=1, color_rgb=RGBColor(62, 0, 117))
        add_paragraph_with_spacing(doc, desc, space_after=8)
    # Save document
    filename = "Zepto_Category_Discovery_Segments_Report_Combined.docx"
    doc.save(filename)
    print(f"Combined segments report saved successfully as '{os.path.abspath(filename)}'")

if __name__ == '__main__':
    create_combined_segments_report()
